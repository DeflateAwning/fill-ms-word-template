#!/usr/bin/env python3

# Fill a Word document from a template

# Functionality:
# * reads a table with a list of docs to create
# * reads a template which is filled with the content from the first row of that table
# * generates a new file for each row in the table, replacing the content from the first row with content from the target row
# * any cols that you don't want to replace, prepend the col name with "_"

import argparse
from random import choice
import time, datetime
import re, sys, os, shutil

from loguru import logger
import easygui as g

import docx

import pandas as pd

logger.add(sys.stderr, format="{time} {level} {message}", filter=__file__, level="INFO")

def do_many_replacements(orig_str: str, replace_dict: dict) -> str:
    """ Within orig_str, replaces all keys with values from replace_dict. """
    for key, val in replace_dict.items():
        if key is None or val is None or pd.isna(key) or pd.isna(val):
            continue
        
        key = str(key)
        val = str(val)

        orig_str = orig_str.replace(key, val)
    return orig_str


def docx_replace_regex(doc_obj: docx.Document, regex, replace: str):
    """ Source: https://stackoverflow.com/a/42829667 """

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

def duplicate_fill_docx(template_path: str, replace_table: pd.DataFrame, out_folder_path: str):
    """ Duplicates template, fills it based on the replace_table (using the first row as keys), and then saves it. """

    ## Get ready for replacement info
    df = replace_table.copy()
    df = df.drop(columns=[col for col in df.columns if col.startswith('_')])
    df = df[pd.notna(df[df.columns[0]])] # drop rows where the first col is NaN
    df.columns = df.iloc[0] # rename columns to the values in the first row, to be used as the replacement keys

    for index, row in df.iterrows():
        logger.info(f"Starting output row number {index}.")

        ## Make replace_dict
        replace_dict = row.to_dict()

        ## Duplicate the file
        new_filename = os.path.basename(template_path)
        new_filename = do_many_replacements(new_filename, replace_dict)
        new_file_path = os.path.join(out_folder_path, new_filename)
        #shutil.copy(template_path, new_file_path)
        logger.info(f'Will copy to new file: {new_filename}')

        ## Do replacement within file
        doc = docx.Document(template_path)
        for key, val in replace_dict.items():
            if key is None or val is None:
                continue
            
            key = str(key)
            val = str(val)

            docx_replace_regex(doc, re.compile(re.escape(key)), val)
        
        doc.save(new_file_path)

        #breakpoint()
        logger.info('Saved file with replacements.')
        


def start_fill():
    """ Prompts user for args to duplicate_fill_docx, then does it. """
    logger.info('Starting start_fill()')
    template_path = g.fileopenbox("Select the template Word document.", filetypes=['*.docx'])

    table_path = g.fileopenbox("Select the table file.", filetypes=['*.xlsx'])
    replace_table = pd.read_excel(table_path)

    out_folder_path = g.diropenbox("Select the output folder.")

    duplicate_fill_docx(template_path, replace_table, out_folder_path)

    logger.info('Done start_fill()')

if __name__ == '__main__':
    start_fill()

